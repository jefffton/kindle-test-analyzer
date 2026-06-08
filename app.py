import os
import re
import json
import csv
import io
import threading
import uuid
import time
import chardet
from flask import Flask, request, render_template, jsonify
from werkzeug.utils import secure_filename
import anthropic

# Document parsers
from docx import Document
from pdfminer.high_level import extract_text as pdf_extract_text
from openpyxl import load_workbook
from bs4 import BeautifulSoup

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')

ALLOWED_EXTENSIONS = {
    'txt', 'doc', 'docx', 'pdf', 'xlsx', 'xls', 'csv',
    'html', 'htm', 'md', 'json', 'xml', 'rtf'
}

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# In-memory job store: {job_id: {status, result, error, created_at}}
_jobs = {}
_jobs_lock = threading.Lock()


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_file(filepath, filename):
    ext = filename.rsplit('.', 1)[1].lower()

    if ext == 'pdf':
        return pdf_extract_text(filepath)

    if ext in ('docx', 'doc'):
        doc = Document(filepath)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return '\n'.join(parts)

    if ext == 'xlsx':
        wb = load_workbook(filepath, read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f'[Sheet: {sheet.title}]')
            for row in sheet.iter_rows(values_only=True):
                row_text = ' | '.join(str(c) for c in row if c is not None)
                if row_text.strip():
                    parts.append(row_text)
        return '\n'.join(parts)

    if ext == 'xls':
        import xlrd

        # Try xlrd first (true legacy .xls binary format)
        try:
            wb = xlrd.open_workbook(filepath)
            parts = []
            for sheet in wb.sheets():
                parts.append(f'[Sheet: {sheet.name}]')
                for row_idx in range(sheet.nrows):
                    row_text = ' | '.join(
                        str(sheet.cell_value(row_idx, c))
                        for c in range(sheet.ncols)
                        if sheet.cell_value(row_idx, c) != ''
                    )
                    if row_text.strip():
                        parts.append(row_text)
            return '\n'.join(parts)
        except Exception:
            pass

        # Fallback: openpyxl (file may be xlsx with wrong extension)
        try:
            wb = load_workbook(filepath, read_only=True, data_only=True)
            parts = []
            for sheet in wb.worksheets:
                parts.append(f'[Sheet: {sheet.title}]')
                for row in sheet.iter_rows(values_only=True):
                    row_text = ' | '.join(str(c) for c in row if c is not None)
                    if row_text.strip():
                        parts.append(row_text)
            return '\n'.join(parts)
        except Exception:
            pass

        # Fallback: HTML (some Excel exports are HTML with .xls extension)
        try:
            raw = open(filepath, 'rb').read()
            enc = chardet.detect(raw)['encoding'] or 'utf-8'
            soup = BeautifulSoup(raw.decode(enc, errors='replace'), 'html.parser')
            text = soup.get_text(separator='\n')
            if text.strip():
                return text
        except Exception:
            pass

        raise ValueError('Could not parse .xls file — try converting it to .xlsx and re-uploading.')

    if ext == 'csv':
        raw = open(filepath, 'rb').read()
        enc = chardet.detect(raw)['encoding'] or 'utf-8'
        text = raw.decode(enc, errors='replace')
        reader = csv.reader(io.StringIO(text))
        return '\n'.join(' | '.join(row) for row in reader if any(row))

    if ext in ('html', 'htm'):
        raw = open(filepath, 'rb').read()
        enc = chardet.detect(raw)['encoding'] or 'utf-8'
        soup = BeautifulSoup(raw.decode(enc, errors='replace'), 'html.parser')
        return soup.get_text(separator='\n')

    if ext == 'json':
        raw = open(filepath, 'rb').read()
        enc = chardet.detect(raw)['encoding'] or 'utf-8'
        data = json.loads(raw.decode(enc, errors='replace'))
        return json.dumps(data, indent=2)

    raw = open(filepath, 'rb').read()
    enc = chardet.detect(raw)['encoding'] or 'utf-8'
    return raw.decode(enc, errors='replace')


def build_prompt(text_content, filename, feature_context='', tester_name=''):
    doc_snippet = text_content[:5000]

    feature_section = ''
    if feature_context and feature_context.strip():
        feature_section = f"""
FEATURE CONTEXT (provided by the tester — use this to sharpen your analysis):
---
{feature_context.strip()}
---
"""

    personalize_section = ''
    if tester_name and tester_name.strip():
        personalize_section = f"""
PERSONALIZATION: The tester's name is "{tester_name.strip()}". Use their name as the main character in EVERY scenario instead of generic personas. Write scenarios in a personal, story-like way — as if {tester_name.strip()} is the one using the Kindle. For example: "{tester_name.strip()} is lying in bed at midnight and decides to buy one more book..." This makes scenarios feel relatable and memorable for the tester.
"""

    return f"""You are a senior exploratory tester for Amazon Kindle Storefront with 10+ years of experience.
You stay current with modern ad-hoc testing trends: chaos engineering for UI, accessibility-first testing,
AI/ML-driven feature testing, privacy/data-handling exploration, dark-pattern detection, cross-app deep linking,
session-resumption testing, low-memory device degradation, network shaping (2G/spotty 5G/airplane mode toggling),
voice assistant integration (Alexa), Bluetooth device hand-offs, and real-world battery/thermal stress scenarios.

File: "{filename}"
{feature_section}{personalize_section}
========================================================================
EXISTING TEST CASES IN THE UPLOADED DOCUMENT (read carefully — your scenarios MUST NOT overlap with these):
========================================================================
{doc_snippet}
========================================================================

ABSOLUTE HARD RULES — READ TWICE:

RULE 1 — ZERO DUPLICATION:
- Read every test case in the uploaded document above.
- Your exploratory scenarios MUST be substantively different — not just reworded versions.
- A scenario is a DUPLICATE if its core happy-path action matches any existing TC's action
  (e.g. "user searches for a book" matches an existing search TC even if you add a different device).
- Before finalizing each scenario, mentally check: "Could this scenario be marked PASS by simply
  running the closest existing TC?" If yes, DISCARD it and write a different one.

RULE 2 — IF YOU CAN'T FIND TRULY NEW IDEAS, RETURN FEWER SCENARIOS:
- Quality > quantity. If you can only generate 5 genuinely novel scenarios, return 5 — not 10 padded ones.
- It is BETTER to return 4-5 sharp, original scenarios than 10 that overlap with the document.
- DO NOT pad output by recycling the uploaded test cases in disguise.

RULE 3 — EXPLORATORY MEANS UNSCRIPTED:
- Exploratory scenarios are NOT step-by-step verification of expected behavior.
- They are investigations driven by curiosity: "What happens if...?", "Can a user accidentally...?",
  "What if two things happen at once?"
- Each scenario must include at least ONE of these qualities:
  * Concurrency / race condition (two actions at the same moment)
  * State pollution (corrupted local cache, partial sync, orphan data)
  * Modern threat model (privacy, dark patterns, accessibility regression, AI hallucination, deep-link hijack)
  * Cross-surface interaction (Kindle + Alexa, Kindle + Goodreads, Kindle + Audible hand-off)
  * Device degradation (low storage, low battery, thermal throttle, screen burn-in on e-ink)
  * Real-world chaos (network flapping, OS-level interruption, app backgrounding mid-action)

RULE 4 — MODERN AD-HOC TESTING TRENDS TO COVER (pick the most relevant 5-7 for this document):
- AI-powered features: AI-generated book summaries, recommendations bias, hallucinated content, training-data leakage
- Privacy & data residency: GDPR right-to-delete, regional content licensing edge cases, kids' COPPA flows
- Dark-pattern detection: forced opt-ins, hidden subscription renewals, tricky "cancel" flows
- Accessibility regressions: VoiceView, dynamic text scaling, high-contrast, switch control, color-blind safe palettes
- Deep-link / universal-link hijack: malicious URLs, expired share links, region-mismatched deep links
- Cross-device session handoff: "continue reading on another device" failures, sync conflicts, last-page disagreements
- Subscription / payment edge cases: trial-to-paid race, gift card + KU combo, expired card mid-borrow, family library sharing limits
- Offline-first behavior: airplane-mode purchase queue, partial sample mid-network-loss, library access without auth refresh
- AI/ML recommendation poisoning: rapid wishlist-then-remove churn, watching how recommendations adapt
- Voice integration: Alexa "read me chapter 3" with ambiguous title, Echo Show display fallback
- Low-end device behavior: 2GB-RAM Fire HD with 50+ open tabs, e-ink ghosting on Paperwhite after long reading
- Cross-app interruptions: phone call mid-purchase, screen recorder running during DRM content, picture-in-picture audiobook
- Localization edge cases: RTL languages (Arabic, Hebrew), CJK fonts on e-ink, currency conversion precision

Return ONLY a JSON object (no markdown fences, no explanation before or after):

{{
  "document_summary": "2-3 sentences describing the uploaded document's scope and what it tests",
  "test_cases_found": [
    {{"id": "TC-001", "title": "title", "category": "category", "description": "what it tests", "priority": "High|Medium|Low"}}
  ],
  "coverage_gaps": ["A specific user behavior or edge case that the existing test cases do NOT cover"],
  "exploratory_scenarios": [
    {{
      "id": "ES-001",
      "title": "short distinctive name",
      "persona": "who is doing this",
      "device": "which device",
      "scenario": "Write as a real user story with 4-7 numbered steps. Make at least one step messy, surprising, or modern (concurrency, AI feature, privacy concern, accessibility need, etc.). Avoid generic happy-path search/buy/read flows that any existing TC already covers.",
      "what_to_look_for": "Concrete observations, modern bug categories (UX, perf, privacy, accessibility, AI behavior, security), and how to confirm if it broke",
      "risk_level": "High|Medium|Low",
      "novelty_check": "ONE sentence stating WHY this scenario is NOT covered by any existing test case in the document",
      "trend_tag": "which modern testing trend this exercises (e.g. AI-recommendations-bias, deep-link-hijack, cross-device-handoff, accessibility-regression, dark-pattern, privacy, chaos-engineering, voice-integration, low-end-device, offline-first)",
      "tags": ["tag1", "tag2"]
    }}
  ],
  "kindle_specific_risks": [
    {{"area": "risk area", "description": "what could go wrong from a real user's perspective", "suggested_test": "a concrete thing to try"}}
  ],
  "total_test_cases": 0,
  "total_scenarios": 0
}}

FINAL CHECK before responding:
- For each exploratory_scenario, the novelty_check field MUST justify why no existing TC covers it.
- If you cannot fill novelty_check honestly for a scenario, DELETE that scenario.
- Aim for 6-10 scenarios. If fewer than 6 are genuinely novel, return only the novel ones.
- At least 4 scenarios must use a "trend_tag" from the modern trends list above.
- Identify at least 5 coverage gaps and 4 Kindle-specific risks (these can repeat themes from existing TCs since their job is to point out what's missing)."""


def call_claude(prompt):
    """Call Claude API. Uses Claude CLI in AgentSpaces, or direct SDK elsewhere."""
    import subprocess

    claude_bin = '/agentspaces/cecelia/claude'
    use_cli = os.path.exists(claude_bin)

    if use_cli:
        # AgentSpaces environment — use authenticated CLI
        model = os.environ.get('ANTHROPIC_DEFAULT_SONNET_MODEL', 'global.anthropic.claude-sonnet-4-6[1m]')
        result = subprocess.run(
            [claude_bin, '--print', '--output-format', 'json', '--model', model, '--bare'],
            input=prompt,
            capture_output=True,
            text=True,
            timeout=600,
            env={**os.environ,
                 'ANTHROPIC_API_KEY': os.environ.get('ANTHROPIC_API_KEY', 'placeholder'),
                 'CLAUDE_CODE_DISABLE_NONESSENTIAL_TRAFFIC': '1'}
        )
        if result.returncode != 0:
            raise RuntimeError(f'Claude CLI error: {result.stderr[:300]}')
        cli_output = json.loads(result.stdout)
        if cli_output.get('is_error'):
            raise RuntimeError(f'Claude error: {cli_output.get("result", "unknown")}')
        return cli_output.get('result', '').strip()
    else:
        # Check which provider to use
        gemini_key = os.environ.get('GEMINI_API_KEY', '')
        anthropic_key = os.environ.get('ANTHROPIC_API_KEY', '')

        if gemini_key:
            # Google Gemini (free tier)
            import google.generativeai as genai
            genai.configure(api_key=gemini_key)
            model = genai.GenerativeModel(os.environ.get('GEMINI_MODEL', 'gemini-2.5-flash'))
            response = model.generate_content(prompt)
            return response.text.strip()
        elif anthropic_key:
            # Anthropic Claude (paid)
            use_bedrock = os.environ.get('USE_BEDROCK', '') == '1'
            if use_bedrock:
                client = anthropic.AnthropicBedrock(
                    aws_access_key=os.environ.get('AWS_ACCESS_KEY_ID'),
                    aws_secret_key=os.environ.get('AWS_SECRET_ACCESS_KEY'),
                    aws_session_token=os.environ.get('AWS_SESSION_TOKEN'),
                    aws_region=os.environ.get('AWS_DEFAULT_REGION', 'us-east-1'),
                )
                model = os.environ.get('BEDROCK_MODEL', 'us.anthropic.claude-sonnet-4-5-20251001-v1:0')
            else:
                client = anthropic.Anthropic(api_key=anthropic_key)
                model = os.environ.get('ANTHROPIC_MODEL', 'claude-sonnet-4-6')

            message = client.messages.create(
                model=model,
                max_tokens=4096,
                messages=[{'role': 'user', 'content': prompt}]
            )
            return message.content[0].text.strip()
        else:
            raise RuntimeError(
                'No API key found. Set GEMINI_API_KEY (free) or ANTHROPIC_API_KEY in your environment.'
            )


def run_analysis(job_id, text_content, filename, feature_context='', tester_name=''):
    with _jobs_lock:
        _jobs[job_id]['status'] = 'running'

    try:
        prompt = build_prompt(text_content, filename, feature_context, tester_name)
        response_text = call_claude(prompt)

        # Extract JSON — handles preamble text + optional markdown fences
        json_match = re.search(r'```(?:json)?\s*(\{[\s\S]+?\})\s*```', response_text)
        if json_match:
            response_text = json_match.group(1)
        else:
            start = response_text.find('{')
            end = response_text.rfind('}')
            if start != -1 and end != -1:
                response_text = response_text[start:end + 1]

        data = json.loads(response_text)
        data['total_test_cases'] = len(data.get('test_cases_found', []))
        data['total_scenarios'] = len(data.get('exploratory_scenarios', []))
        data['filename'] = filename

        with _jobs_lock:
            _jobs[job_id]['status'] = 'done'
            _jobs[job_id]['result'] = data

    except Exception as e:
        with _jobs_lock:
            _jobs[job_id]['status'] = 'error'
            _jobs[job_id]['error'] = str(e)


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/analyze', methods=['POST'])
def analyze():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if not file.filename:
        return jsonify({'error': 'No file selected'}), 400

    if not allowed_file(file.filename):
        return jsonify({
            'error': f'File type not supported. Allowed: {", ".join(sorted(ALLOWED_EXTENSIONS))}'
        }), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    try:
        text = extract_text_from_file(filepath, filename)
        if not text or not text.strip():
            return jsonify({'error': 'Could not extract text from the file. Is it empty or image-only?'}), 400
    except Exception as e:
        return jsonify({'error': f'File parsing error: {str(e)}'}), 400
    finally:
        try:
            os.remove(filepath)
        except OSError:
            pass

    feature_context = request.form.get('feature_context', '').strip()
    tester_name = request.form.get('tester_name', '').strip()

    job_id = str(uuid.uuid4())
    with _jobs_lock:
        _jobs[job_id] = {
            'status': 'queued',
            'result': None,
            'error': None,
            'created_at': time.time()
        }

    t = threading.Thread(target=run_analysis, args=(job_id, text, filename, feature_context, tester_name), daemon=True)
    t.start()

    return jsonify({'job_id': job_id})


@app.route('/result/<job_id>')
def get_result(job_id):
    with _jobs_lock:
        job = _jobs.get(job_id)

    if not job:
        return jsonify({'error': 'Job not found'}), 404

    if job['status'] in ('queued', 'running'):
        return jsonify({'status': job['status']})

    if job['status'] == 'error':
        return jsonify({'status': 'error', 'error': job['error']}), 500

    return jsonify({'status': 'done', 'data': job['result']})


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port, debug=False, threaded=True)
